import hashlib
import os
import uuid

import docx
import duckdb
from pypdf import PdfReader
from rq import get_current_job
from rq.job import Job

import app.main
from app import logging_config
from app.main import vector_db_client, embedding_client
from app.models import Document, ParquetSchema
from app.vector_db_clients import VectorPoint
from app import database

logger = logging_config.setup_logging("nexus-ingestion-worker")


def get_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text() + "\n"
    return text


def get_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def chunk_text(text, chunk_size=512, overlap=51):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunks.append(" ".join(words[i:i + chunk_size]))
    return chunks


def _get_batch_embeddings(chunks: list[str]) -> list[list[float]]:
    if not chunks:
         return []
    try:
        return embedding_client.embed(chunks, timeout=30.0)
    except Exception as e:
        logger.error(f"Failed to batch embed: {e}", exc_info=True)
        # Return empty list to signal failure, or zero vectors
        return [[0.0]*384 for _ in chunks]


def _update_batch_status(job: Job, status_key: str):
    """
    Updates the Redis scoreboard.
    status_key should be 'completed' or 'failed'.
    """
    if batch_id := job.meta.get("batch_id"):
        conn = job.connection
        conn.hincrby(f"batch:{batch_id}", status_key, 1)

        data = conn.hgetall(f"batch:{batch_id}")
        total = int(data.get(b'total', 0))
        done = int(data.get(b'completed', 0)) + int(data.get(b'failed', 0))

        if done >= total:
            conn.hset(f"batch:{batch_id}", "status", "finished")


def _get_file_hash(file_path: str) -> str:
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


def _get_catalog_entry(db_session, file_hash: str):
    """Checks if the hash exists in the SQL documents table."""
    return db_session.query(Document).filter(Document.content_hash == file_hash).first()


def _create_catalog_entry(db_session, file_name: str, file_path: str, file_hash: str, status: str):
    """Records a new file in the SQL documents table."""
    file_size = os.path.getsize(file_path)
    document = Document(
        file_name=file_name,
        file_path=file_path,
        content_hash=file_hash,
        file_size=file_size,
        status=status
    )
    db_session.add(document)
    db_session.commit()


def extract_schema(filepath: str) -> str:
    try:
        conn = duckdb.connect(":memory:")
        res = conn.execute(f"DESCRIBE SELECT * FROM '{filepath}'").fetchall()
        cols = [f"{row[0]} ({row[1]})" for row in res]
        return ", ".join(cols)
    except Exception as e:
        logger.error(f"Error extracting schema from {filepath}: {e}", exc_info=True)
        raise


def _handle_parquet_file(db_session, file_path: str, job):
    try:
        schema_str = extract_schema(file_path)
        table_name = os.path.basename(file_path).replace(".parquet", "")
        existing_table = db_session.query(ParquetSchema).filter(
            ParquetSchema.table_name == table_name
        ).first()

        # Update or Create
        if existing_table:
            if existing_table.columns != schema_str:
                existing_table.columns = schema_str
                db_session.commit()
                logger.info(f"Updated schema for {table_name}")
            else:
                logger.info(f"Duplicate detected - same schema. Skipping")
                _update_batch_status(job, "completed")
                return {"status": "skipped", "reason": "duplicate_schema", "file": file_path}
        else:
            new_schema = ParquetSchema(
                table_name=table_name,
                file_path=file_path,
                columns=schema_str
            )
            db_session.add(new_schema)
            db_session.commit()
            logger.info(f"Added new schema for {table_name}")
    except Exception as e:
        logger.error(f"Failed to handle parquet file: {e}", exc_info=True)
        _update_batch_status(job, "failed")

    _update_batch_status(job, "completed")
    return {"status": "success", "file": file_path}


def _handle_document(db_session, file_path: str, job):
    file_name = os.path.basename(file_path)

    # 1. IMMEDIATE HASH CHECK (The Idempotency Gate)
    try:
        file_hash = _get_file_hash(file_path)
        existing_record = _get_catalog_entry(db_session, file_hash)

        if existing_record:
            logger.info(f"Duplicate detected (Hash: {file_hash}). Skipping embedding.")
            # Even if we skip, we MUST update the batch status so the progress bar finishes
            _update_batch_status(job, "completed")
            return {"status": "skipped", "reason": "duplicate_hash", "file": file_path}

    except Exception as e:
        logger.error(f"Failed to perform hash check: {e}", exc_info=True)
        _update_batch_status(job, "failed")
        return {"status": "error", "message": "Hash calculation failed"}

    # 2. Extract Text
    try:
        if file_path.endswith('.pdf'):
            text = get_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            text = get_text_from_docx(file_path)
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            logger.warning(f"Skipping unsupported file type: {file_path}")
            _update_batch_status(job, "failed")
            return {"status": "skipped", "reason": "unsupported_extension"}

        chunks = chunk_text(text)
        if not chunks:
            _update_batch_status(job, "completed")
            return {"status": "success", "file": file_path, "chunks": 0}

        # 3. Embed & Vector Upsert
        try:
            collection_name = app.main.document_retriever.collection_name
            batch_size = 32
            docs_inserted = 0

            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                embeddings = _get_batch_embeddings(batch)
                points = []

                for j, (chunk, emb) in enumerate(zip(batch, embeddings)):
                    chunk_idx = i + j
                    # Deterministic ID based on Hash + Index prevents vector duplication
                    point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{file_hash}_{chunk_idx}"))
                    points.append(
                        VectorPoint(
                            id=point_id,
                            vector=emb,
                            payload={"filename": file_name, "chunk": chunk_idx, "text": chunk}
                        )
                    )

                vector_db_client.upsert(collection_name=collection_name, points=points)
                docs_inserted += len(points)
        except Exception as e:
            _create_catalog_entry(db_session, file_name, file_path, file_hash, 'FAILED')
            logger.error(f"Worker Error during document embedding: {e}", exc_info=True)
            return {"status": "failed", "reason": "embedding_error"}

        # 4. Success: Update Catalog and Batch Status
        _create_catalog_entry(db_session, file_name, file_path, file_hash, 'INGESTED')
        _update_batch_status(job, "completed")
        return {"status": "success", "file": file_path, "chunks": docs_inserted}
    except Exception as e:
        logger.error(f"Worker Error during processing: {e}", exc_info=True)
        _update_batch_status(job, "failed")
        return {"status": "error", "file": file_path}


def process_file_job(file_path: str):
    logger.info(f"Worker Processing: {file_path}")
    job = get_current_job()
    assert job is not None, "Job should not be None"

    with database.get_db_ctx() as db_session:
        if file_path.endswith('.parquet'):
            result = _handle_parquet_file(db_session, file_path, job)
        else:
            result = _handle_document(db_session, file_path, job)

        job.meta['progress_percentage'] = 100
        job.save_meta()

        return result